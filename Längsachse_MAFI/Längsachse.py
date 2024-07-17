from docx import Document
from docx.shared import Inches, Cm, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
from docx2pdf import convert
import os

def create_word_document(measuring_distances, degrees, Titel, name):
    # Create a new Word document
    doc = Document()

    # Add a new header
    section = doc.sections[0]
    header = section.header
    header_paragraph = header.paragraphs[0]
    header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run = header_paragraph.add_run()
    run.add_picture(r'C:\Users\daka\Documents\Längsachse_MAFI\cedes_logo_mit_claim_rgb_1.png', width=Inches(1.0))

    # Add a heading with a specific style using the style name as a key
    heading_title = f"{Titel}"
    heading = doc.add_heading(heading_title, level=1)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add an empty paragraph to create a gap between the title and the table
    gap_paragraph = doc.add_paragraph()
    gap_paragraph.add_run().add_break()  # Adds a line break
    gap_paragraph_format = gap_paragraph.paragraph_format
    gap_paragraph_format.space_after = Pt(24)  # Sets the space after the paragraph to 24 points

    # Add footer with date and name
    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    footer_paragraph.text = f'Date: {datetime.now().strftime("%Y-%m-%d")}                                                      		Name: {name}'

    # Add a table with measuring distances and degrees
    table = doc.add_table(rows=len(measuring_distances) + 1, cols=2)
    table.style = 'Table Grid'  # Add gridlines to the table

    # Set the column width to 4 cm
    widths = [Cm(4), Cm(4)]
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

    # Fill the table headers
    table.cell(0, 0).text = 'Distance [cm]'
    table.cell(0, 1).text = 'Degree [°]'
    table.cell(0, 0).paragraphs[0].runs[0].bold = True
    table.cell(0, 1).paragraphs[0].runs[0].bold = True

    # Fill the table with measurements
    for i, (distance, degree) in enumerate(zip(measuring_distances, degrees)):
        table.cell(i + 1, 0).text = f'{distance} cm'
        table.cell(i + 1, 1).text = f'{degree}°'

    # Convert Word document to PDF
    folder_path = "PDF_Documents"
    os.makedirs(folder_path, exist_ok=True)  # Create folder if not exists
    current_datetime = datetime.now().strftime("%Y%m%d_%H%M")
    pdf_file_path = os.path.join(folder_path, f'document_{current_datetime}_{name}.pdf')
    docx_file_path = os.path.join(folder_path, f'document_{current_datetime}_{name}.docx')
    doc.save(docx_file_path)

    # Convert Word document to PDF
    convert(docx_file_path, pdf_file_path)

    # Optionally, delete the intermediate DOCX file
    os.remove(docx_file_path)

if __name__ == "__main__":
    # Example usage with measuring_distances, Titel, and name
    measuring_distances = [1, 2, 3, 4, 5]
    degrees = [1, 2, 3, 4, 5]
    Titel = 'Cegard Pro Gradmessungen'
    name = 'Your Name'
    create_word_document(measuring_distances, degrees, Titel, name)
